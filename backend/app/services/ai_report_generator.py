"""
AI Report Generator — Gemini-powered narrative reports as .docx downloads.

Service layer over the AI features. All four report types share the same
centralised Gemini AIClient (configured via settings.GEMINI_API_KEY).

Public functions:
    generate_lead_journey_report(timeline_data, lead_name) -> str
    generate_periodic_leads_report(summary, period_label, target_name) -> str
    generate_user_performance_report(metrics, period_label, user_name, user_role) -> str
    generate_team_performance_report(metrics) -> str
    build_docx_report(title, body_text) -> BytesIO
"""

import io
import logging

import matplotlib
matplotlib.use('Agg')  # non-interactive backend, safe for async servers
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.ai.client import get_ai_client
from app.ai.config import ai_settings
from app.ai.features.report_generation import (
    LeadJourneyFeature,
    PeriodicLeadsReportFeature,
    UserPerformanceReportFeature,
    TeamPerformanceFeature,
)
from app.ai.exceptions import AIServiceError

logger = logging.getLogger(__name__)


class AIReportError(Exception):
    """Raised when the AI report generation fails (e.g. Gemini API error).

    Re-exported at module level so that reports.py's
    `from app.services.ai_report_generator import AIReportError`
    continues to work without modification.
    """
    pass


async def generate_lead_journey_report(timeline_data: list[dict], lead_name: str) -> str:
    """Send timeline events to Gemini and get a narrative summary."""
    client = get_ai_client()
    feature = LeadJourneyFeature(client=client)
    try:
        result = await feature.run(
            {"timeline_data": timeline_data, "lead_name": lead_name},
            entity_id="lead_journey",
            model_name=ai_settings.AI_REPORT_MODEL_NAME,
        )
        return result.text
    except AIServiceError as exc:
        logger.exception("Gemini API call failed for lead journey report")
        raise AIReportError(
            f"AI report generation failed for lead '{lead_name}'"
        ) from exc


async def generate_periodic_leads_report(
    summary: dict,
    period_label: str,
    target_name: str,
) -> str:
    """Generate a combined portfolio overview + lead journey narrative for a period."""
    client = get_ai_client()
    feature = PeriodicLeadsReportFeature(client=client)
    try:
        result = await feature.run(
            {"summary": summary, "period_label": period_label, "target_name": target_name},
            entity_id="periodic_leads",
            model_name=ai_settings.AI_REPORT_MODEL_NAME,
        )
        return result.text
    except AIServiceError as exc:
        logger.exception("Gemini API call failed for periodic leads report")
        raise AIReportError(
            f"AI report generation failed for periodic leads report"
        ) from exc


async def generate_user_performance_report(
    metrics: dict,
    period_label: str,
    user_name: str,
    user_role: str,
) -> str:
    """Generate an individual performance review narrative."""
    client = get_ai_client()
    feature = UserPerformanceReportFeature(client=client)
    try:
        result = await feature.run(
            {
                "metrics": metrics,
                "period_label": period_label,
                "user_name": user_name,
                "user_role": user_role,
            },
            entity_id="user_performance",
            model_name=ai_settings.AI_REPORT_MODEL_NAME,
        )
        return result.text
    except AIServiceError as exc:
        logger.exception("Gemini API call failed for user performance report")
        raise AIReportError(
            f"AI report generation failed for user performance report"
        ) from exc


async def generate_team_performance_report(metrics: dict) -> str:
    """Send team metrics to Gemini and get a performance digest."""
    client = get_ai_client()
    feature = TeamPerformanceFeature(client=client)
    try:
        result = await feature.run(
            {"metrics": metrics},
            entity_id="team_performance",
            model_name=ai_settings.AI_REPORT_MODEL_NAME,
        )
        return result.text
    except AIServiceError as exc:
        logger.exception("Gemini API call failed for team performance report")
        raise AIReportError(
            "AI report generation failed for team performance report"
        ) from exc



# ── Chart generators ──────────────────────────────────────────────────────────

STATUS_COLORS = {
    "new": "#3b82f6", "in_progress": "#f59e0b", "potential": "#8b5cf6",
    "converted_to_investor": "#10b981", "existing_investor": "#059669",
    "non_potential": "#ef4444",
}
STATUS_LABELS = {
    "new": "New", "in_progress": "In Progress", "potential": "Potential",
    "converted_to_investor": "Converted", "existing_investor": "Existing Investor",
    "non_potential": "Non-Potential",
}

# Brand palette
BRAND_BLUE = "#053abb"
BRAND_GREEN = "#10b981"
BRAND_AMBER = "#f59e0b"
BRAND_PURPLE = "#8b5cf6"
BRAND_RED = "#ef4444"
CHART_PALETTE = [BRAND_BLUE, BRAND_GREEN, BRAND_AMBER, BRAND_PURPLE, BRAND_RED, "#0ea5e9", "#ec4899"]


def _fig_to_png(fig) -> io.BytesIO:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=120, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_pipeline_distribution(by_status: dict) -> io.BytesIO | None:
    if not by_status:
        return None
    labels = [STATUS_LABELS.get(k, k) for k in by_status]
    values = list(by_status.values())
    colors = [STATUS_COLORS.get(k, "#94a3b8") for k in by_status]
    fig, ax = plt.subplots(figsize=(6, 3.5))
    ax.pie(values, labels=labels, colors=colors, autopct="%1.0f%%",
           startangle=90, pctdistance=0.75,
           wedgeprops={"edgecolor": "white", "linewidth": 1.5})
    ax.set_title("Pipeline Distribution", fontweight="bold", color="#1e3a8a", pad=12)
    fig.patch.set_facecolor("#f8fafc")
    return _fig_to_png(fig)


def _chart_by_source(by_source: dict) -> io.BytesIO | None:
    if not by_source:
        return None
    items = sorted(by_source.items(), key=lambda x: x[1], reverse=True)[:8]
    labels, values = zip(*items)
    fig, ax = plt.subplots(figsize=(6, 3.5))
    bars = ax.barh(labels, values, color="#1d4ed8", edgecolor="white")
    ax.set_xlabel("Number of Leads", color="#374151")
    ax.set_title("Leads by Acquisition Channel", fontweight="bold", color="#1e3a8a", pad=10)
    ax.tick_params(axis="both", labelsize=8, colors="#374151")
    ax.spines[["top", "right"]].set_visible(False)
    for bar, val in zip(bars, values):
        ax.text(val + 0.05, bar.get_y() + bar.get_height() / 2,
                str(val), va="center", fontsize=8, color="#374151")
    fig.patch.set_facecolor("#f8fafc")
    fig.tight_layout()
    return _fig_to_png(fig)


def _chart_performance_bars(metrics: dict) -> io.BytesIO | None:
    by_status = metrics.get("by_status", {})
    if not by_status:
        return None
    labels = [STATUS_LABELS.get(k, k) for k in by_status]
    values = list(by_status.values())
    colors = [STATUS_COLORS.get(k, "#94a3b8") for k in by_status]
    fig, ax = plt.subplots(figsize=(6, 3.2))
    ax.bar(labels, values, color=colors, edgecolor="white")
    ax.set_ylabel("Count", color="#374151")
    ax.set_title("Lead Pipeline Breakdown", fontweight="bold", color="#1e3a8a", pad=10)
    ax.tick_params(axis="x", labelrotation=20, labelsize=7.5)
    ax.spines[["top", "right"]].set_visible(False)
    fig.patch.set_facecolor("#f8fafc")
    fig.tight_layout()
    return _fig_to_png(fig)


def _chart_team_comparison(members: list) -> io.BytesIO | None:
    if not members:
        return None
    names = [m.get("user_name", "?")[:12] for m in members]
    leads = [m.get("total_leads_assigned", 0) for m in members]
    converted = [m.get("converted_leads", 0) for m in members]
    x = range(len(names))
    fig, ax = plt.subplots(figsize=(7, 3.5))
    w = 0.35
    ax.bar([i - w/2 for i in x], leads, width=w, label="Assigned", color="#1d4ed8")
    ax.bar([i + w/2 for i in x], converted, width=w, label="Converted", color="#10b981")
    ax.set_xticks(list(x))
    ax.set_xticklabels(names, rotation=15, fontsize=8)
    ax.set_ylabel("Count", color="#374151")
    ax.set_title("Team Lead Performance", fontweight="bold", color="#1e3a8a", pad=10)
    ax.legend(fontsize=8)
    ax.spines[["top", "right"]].set_visible(False)
    fig.patch.set_facecolor("#f8fafc")
    fig.tight_layout()
    return _fig_to_png(fig)


def _chart_interaction_timeline(events_over_time: list) -> io.BytesIO | None:
    """Line chart: interaction events per day for a lead's journey."""
    if not events_over_time:
        return None
    dates = [e["date"] for e in events_over_time]
    values = [e["events"] for e in events_over_time]
    fig, ax = plt.subplots(figsize=(7, 3))
    ax.plot(dates, values, marker="o", color=BRAND_BLUE, linewidth=2.2, markersize=5, markerfacecolor="white", markeredgewidth=2)
    ax.fill_between(dates, values, alpha=0.12, color=BRAND_BLUE)
    ax.set_xlabel("Date", color="#374151", fontsize=8)
    ax.set_ylabel("Events", color="#374151", fontsize=8)
    ax.set_title("Interaction Activity Over Time", fontweight="bold", color="#1e3a8a", pad=10)
    ax.tick_params(axis="x", labelrotation=30, labelsize=7, colors="#374151")
    ax.tick_params(axis="y", labelsize=7, colors="#374151")
    ax.spines[["top", "right"]].set_visible(False)
    ax.set_ylim(bottom=0)
    fig.patch.set_facecolor("#f8fafc")
    fig.tight_layout()
    return _fig_to_png(fig)


def _chart_leads_over_time(leads_over_time: list) -> io.BytesIO | None:
    """Dual-line / area chart: leads created vs converted over time."""
    if not leads_over_time:
        return None
    dates = [e["date"] for e in leads_over_time]
    created = [e["created"] for e in leads_over_time]
    converted = [e["converted"] for e in leads_over_time]
    fig, ax = plt.subplots(figsize=(7, 3.2))
    ax.fill_between(dates, created, alpha=0.15, color=BRAND_BLUE)
    ax.fill_between(dates, converted, alpha=0.18, color=BRAND_GREEN)
    ax.plot(dates, created, marker="o", color=BRAND_BLUE, linewidth=2.2, markersize=5,
            markerfacecolor="white", markeredgewidth=2, label="Created")
    ax.plot(dates, converted, marker="s", color=BRAND_GREEN, linewidth=2.2, markersize=5,
            markerfacecolor="white", markeredgewidth=2, label="Converted")
    ax.set_xlabel("Date", color="#374151", fontsize=8)
    ax.set_ylabel("Leads", color="#374151", fontsize=8)
    ax.set_title("Leads Created vs Converted Over Time", fontweight="bold", color="#1e3a8a", pad=10)
    ax.tick_params(axis="x", labelrotation=30, labelsize=7, colors="#374151")
    ax.tick_params(axis="y", labelsize=7, colors="#374151")
    ax.legend(fontsize=8, framealpha=0)
    ax.spines[["top", "right"]].set_visible(False)
    ax.set_ylim(bottom=0)
    fig.patch.set_facecolor("#f8fafc")
    fig.tight_layout()
    return _fig_to_png(fig)


def _chart_task_completion_trend(tasks_over_time: list) -> io.BytesIO | None:
    """Line chart: tasks assigned vs completed per day."""
    if not tasks_over_time:
        return None
    dates = [e["date"] for e in tasks_over_time]
    assigned = [e["assigned"] for e in tasks_over_time]
    completed = [e["completed"] for e in tasks_over_time]
    fig, ax = plt.subplots(figsize=(7, 3))
    ax.fill_between(dates, assigned, alpha=0.12, color=BRAND_BLUE)
    ax.fill_between(dates, completed, alpha=0.18, color=BRAND_GREEN)
    ax.plot(dates, assigned, marker="o", color=BRAND_BLUE, linewidth=2.2, markersize=5,
            markerfacecolor="white", markeredgewidth=2, label="Assigned")
    ax.plot(dates, completed, marker="s", color=BRAND_GREEN, linewidth=2.2, markersize=5,
            markerfacecolor="white", markeredgewidth=2, label="Completed")
    ax.set_xlabel("Date", color="#374151", fontsize=8)
    ax.set_ylabel("Tasks", color="#374151", fontsize=8)
    ax.set_title("Task Assignment & Completion Trend", fontweight="bold", color="#1e3a8a", pad=10)
    ax.tick_params(axis="x", labelrotation=30, labelsize=7, colors="#374151")
    ax.tick_params(axis="y", labelsize=7, colors="#374151")
    ax.legend(fontsize=8, framealpha=0)
    ax.spines[["top", "right"]].set_visible(False)
    ax.set_ylim(bottom=0)
    fig.patch.set_facecolor("#f8fafc")
    fig.tight_layout()
    return _fig_to_png(fig)


def _chart_team_status_stacked(member_status_chart: list) -> io.BytesIO | None:
    """Stacked bar chart: per-member lead status breakdown."""
    if not member_status_chart:
        return None
    statuses = ["new", "in_progress", "potential", "converted_to_investor",
                "existing_investor", "non_potential"]
    names = [row["name"] for row in member_status_chart]
    bottom_vals = [0] * len(names)
    fig, ax = plt.subplots(figsize=(max(7, len(names) * 1.4), 3.8))
    for status in statuses:
        vals = [row.get(status, 0) for row in member_status_chart]
        if sum(vals) == 0:
            continue
        color = STATUS_COLORS.get(status, "#94a3b8")
        ax.bar(names, vals, bottom=bottom_vals, label=STATUS_LABELS.get(status, status), color=color, edgecolor="white")
        bottom_vals = [b + v for b, v in zip(bottom_vals, vals)]
    ax.set_ylabel("Leads", color="#374151", fontsize=8)
    ax.set_title("Team Lead Status Distribution per Member", fontweight="bold", color="#1e3a8a", pad=10)
    ax.tick_params(axis="x", labelrotation=15, labelsize=7.5, colors="#374151")
    ax.tick_params(axis="y", labelsize=7, colors="#374151")
    ax.legend(fontsize=7, loc="upper right", framealpha=0.8)
    ax.spines[["top", "right"]].set_visible(False)
    fig.patch.set_facecolor("#f8fafc")
    fig.tight_layout()
    return _fig_to_png(fig)


def _embed_image(doc: Document, img_buf: io.BytesIO, caption: str = ""):
    """Add a chart image to the DOCX document."""
    doc.add_picture(img_buf, width=Inches(5.5))
    if caption:
        p = doc.add_paragraph(caption)
        p.paragraph_format.space_after = Pt(8)
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def build_docx_report(
    title: str,
    body_text: str,
    metrics: dict | None = None,
    report_type: str | None = None,
) -> io.BytesIO:
    """Render a styled .docx report with embedded charts and AI narrative.

    Args:
        title:       Report title (Heading 1).
        body_text:   AI-generated narrative.
        metrics:     Raw metrics dict for chart generation.
        report_type: One of 'lead_journey', 'periodic_leads',
                     'user_performance', 'team_performance'.
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Header
    hp = doc.add_paragraph()
    hp.paragraph_format.space_after = Pt(2)
    hr = hp.add_run("IWS Finserv — Wealth Management & Advisory")
    hr.font.size = Pt(9)
    hr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    heading = doc.add_heading(title, level=1)
    heading.paragraph_format.space_after = Pt(12)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    doc.add_paragraph()

    # ── Embed charts based on report type ─────────────────────────────
    if metrics and report_type:
        charts_added = False
        if report_type == "periodic_leads":
            # Chart 1: Pipeline distribution (pie)
            img1 = _chart_pipeline_distribution(metrics.get("by_status", {}))
            if img1:
                _embed_image(doc, img1, "Figure 1: Lead Pipeline Distribution")
                charts_added = True
            # Chart 2: Leads by channel (horizontal bar)
            img2 = _chart_by_source(metrics.get("by_source", {}))
            if img2:
                _embed_image(doc, img2, "Figure 2: Leads by Acquisition Channel")
                charts_added = True
            # Chart 3: Leads created vs converted over time (line/area)
            img3 = _chart_leads_over_time(metrics.get("leads_over_time", []))
            if img3:
                _embed_image(doc, img3, "Figure 3: Leads Created vs Converted Over Time")
                charts_added = True

        elif report_type == "user_performance":
            # Chart 1: Pipeline breakdown (bar)
            img1 = _chart_performance_bars(metrics)
            if img1:
                _embed_image(doc, img1, "Figure 1: Lead Pipeline Breakdown")
                charts_added = True
            # Chart 2: Task completion trend (line)
            img2 = _chart_task_completion_trend(metrics.get("tasks_over_time", []))
            if img2:
                _embed_image(doc, img2, "Figure 2: Task Assignment & Completion Trend")
                charts_added = True

        elif report_type == "team_performance":
            # Chart 1: Team leads assigned vs converted (grouped bar)
            members = metrics.get("members", [])
            img1 = _chart_team_comparison(members)
            if img1:
                _embed_image(doc, img1, "Figure 1: Team Lead Performance")
                charts_added = True
            # Chart 2: Per-member status stacked bar
            member_status_chart = metrics.get("member_status_chart", [])
            img2 = _chart_team_status_stacked(member_status_chart)
            if img2:
                _embed_image(doc, img2, "Figure 2: Team Lead Status Distribution per Member")
                charts_added = True

        elif report_type == "lead_journey":
            # Chart 1: Event type bar chart (horizontal)
            by_type = metrics.get("by_event_type", {})
            if by_type:
                labels = list(by_type.keys())
                values = list(by_type.values())
                fig, ax = plt.subplots(figsize=(6, 2.8))
                ax.barh(labels, values, color=BRAND_BLUE, edgecolor="white")
                ax.set_xlabel("Count", color="#374151")
                ax.set_title("Interaction Events by Type", fontweight="bold",
                             color="#1e3a8a", pad=10)
                ax.spines[["top", "right"]].set_visible(False)
                ax.tick_params(labelsize=8, colors="#374151")
                fig.patch.set_facecolor("#f8fafc")
                fig.tight_layout()
                img = _fig_to_png(fig)
                _embed_image(doc, img, "Figure 1: Interaction Events by Type")
                charts_added = True
            # Chart 2: Interaction activity over time (line)
            img2 = _chart_interaction_timeline(metrics.get("events_over_time", []))
            if img2:
                _embed_image(doc, img2, "Figure 2: Interaction Activity Over Time")
                charts_added = True

        if charts_added:
            doc.add_paragraph()

    # ── AI Narrative body ─────────────────────────────────────────────
    for paragraph in body_text.strip().split("\n\n"):
        cleaned = paragraph.strip()
        if not cleaned:
            continue
        if cleaned.isupper() or (len(cleaned) < 80 and cleaned.endswith(":")):
            p = doc.add_heading(cleaned.rstrip(":"), level=2)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        else:
            p = doc.add_paragraph(cleaned)
            p.paragraph_format.space_after = Pt(6)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    run = footer.add_run(
        "This report was generated by IWS Finserv AI Analytics. "
        "It is intended for internal advisory use only."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

