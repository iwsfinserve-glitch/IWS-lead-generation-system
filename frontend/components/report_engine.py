"""
Report generation utilities: chart rendering and DOCX export.

These functions are designed to work entirely in-memory using io.BytesIO
buffers — no temporary files are written to disk at any point.
"""

import io
import matplotlib
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# WHY "Agg": Matplotlib normally tries to open a GUI window to display
# charts. The "Agg" backend tells it "I only want to render to images
# in memory, don't try to open any windows." This is REQUIRED in
# server environments like Streamlit where there is no display.
matplotlib.use("Agg")


def generate_chart_buffer(
    labels: list[str],
    values: list[int | float],
    title: str = "",
    ylabel: str = "",
) -> io.BytesIO:
    """Generate a minimalist monochrome bar chart and return it as an
    in-memory BytesIO buffer (PNG format).

    Args:
        labels: X-axis category labels (e.g. ["Jan", "Feb", "Mar"])
        values: Corresponding numeric values for each bar
        title:  Optional chart title
        ylabel: Optional Y-axis label

    Returns:
        A BytesIO buffer containing the PNG image, rewound to position 0
        and ready to be read by python-docx's add_picture() or any other
        consumer.
    """
    # Create a figure and axis. figsize=(7, 3.5) gives a wide, compact
    # chart that fits nicely in a Word document column.
    fig, ax = plt.subplots(figsize=(7, 3.5))

    # Draw horizontal bars — monochrome dark grey (#333) to match
    # your app's minimalist aesthetic.
    bars = ax.bar(labels, values, color="#333333", edgecolor="#222222", width=0.55)

    # Add value labels on top of each bar for readability in print
    for bar, val in zip(bars, values):
        ax.text(
            bar.get_x() + bar.get_width() / 2,  # center horizontally
            bar.get_height() + 0.5,               # slight offset above bar
            str(val),
            ha="center", va="bottom",
            fontsize=9, color="#333333",
        )

    # Minimalist styling: remove top/right borders ("spines")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("#cccccc")
    ax.spines["bottom"].set_color("#cccccc")

    # Light grey grid lines on Y-axis only — subtle, not distracting
    ax.yaxis.grid(True, color="#e0e0e0", linestyle="--", linewidth=0.5)
    ax.set_axisbelow(True)  # grid lines behind bars, not on top

    if title:
        ax.set_title(title, fontsize=12, fontweight="bold", color="#333333", pad=12)
    if ylabel:
        ax.set_ylabel(ylabel, fontsize=10, color="#666666")

    ax.tick_params(colors="#666666", labelsize=9)
    fig.tight_layout()

    # ── The critical part: save to an in-memory buffer ──
    # Instead of plt.savefig("chart.png") which writes to disk,
    # we pass a BytesIO object. Matplotlib writes the PNG bytes
    # directly into this RAM buffer.
    buffer = io.BytesIO()
    fig.savefig(buffer, format="png", dpi=150, bbox_inches="tight")

    # IMPORTANT: Close the figure to free memory. Without this,
    # Matplotlib keeps every figure in memory, which leaks RAM
    # in a long-running Streamlit server.
    plt.close(fig)

    # ── seek(0): Rewind the buffer head to the start ──
    # After savefig(), the buffer's internal cursor is at the END
    # of the data. Any subsequent .read() would return b"" (empty).
    # seek(0) moves the cursor back to byte 0 so the next consumer
    # (python-docx's add_picture) reads the full PNG from the start.
    buffer.seek(0)

    return buffer


def generate_docx_report(
    title: str,
    summary_text: str,
    chart_buffers: list[tuple[str, io.BytesIO]] | None = None,
) -> bytes:
    """Build a .docx Word document with a heading, text summary, and
    optionally one or more embedded chart images.

    Args:
        title:         The document heading (e.g. "Lead Journey Report")
        summary_text:  The AI-generated text body. Markdown bold (**) and
                       italic (_) markers are stripped for clean Word output.
        chart_buffers:  Optional list of (chart_title, BytesIO_buffer) tuples.
                       Each buffer should contain a PNG image (as returned by
                       generate_chart_buffer). Each chart gets a sub-heading
                       and is inserted at 5.5 inches wide.

    Returns:
        The complete .docx file as raw bytes, ready to be passed directly
        to Streamlit's st.download_button(data=...).

    The two-buffer pattern:
        Buffer 1 (chart_buffer) — PNG image bytes from Matplotlib
        Buffer 2 (doc_buffer)   — the final .docx file bytes
        Both live in RAM. Nothing touches disk.
    """
    doc = Document()

    # ── Heading ──
    doc.add_heading(title, level=1)

    # ── Summary paragraphs ──
    # Strip markdown formatting since Word doesn't render it.
    # Split on double newlines to preserve paragraph breaks.
    clean_text = summary_text.replace("**", "").replace("_", "")
    for paragraph in clean_text.split("\n\n"):
        stripped = paragraph.strip()
        if stripped:
            doc.add_paragraph(stripped)

    # ── Embedded charts ──
    if chart_buffers:
        for chart_title, buffer in chart_buffers:
            doc.add_heading(chart_title, level=2)

            # WHY seek(0) again: The buffer may have already been read
            # once (e.g. for a Streamlit preview). After being read,
            # the cursor is at the end. We must rewind before python-docx
            # reads it, otherwise add_picture gets zero bytes.
            buffer.seek(0)

            # Inches(5.5) makes the chart span most of a standard
            # 8.5" page width with comfortable margins.
            doc.add_picture(buffer, width=Inches(5.5))

    # ── Save the finished document to a second BytesIO buffer ──
    # This is the "two-buffer" pattern:
    #   Buffer 1: chart PNG → fed into doc.add_picture()
    #   Buffer 2: complete .docx → fed into st.download_button()
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)  # Rewind so the caller can read from the start

    return doc_buffer.getvalue()
